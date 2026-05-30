"""生成约 16000 字的毕业设计说明书 DOCX。

本脚本直接使用用户指定的论文模板：
`模板-基于深度学习的废弃物数据管理系统毕业设计说明书.docx`。

处理方式：
1. 读取模板中的页面设置和基础样式；
2. 清空原有旅游路线规划系统正文；
3. 写入“基于深度学习的废弃物数据管理系统”论文扩写版内容；
4. 在需要后期补图的位置保留醒目的截图占位说明。

脚本只重写目标 DOCX 的正文内容，不删除任何项目文件。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from generate_thesis_docx import (
    add_center,
    add_formula,
    add_heading,
    add_paragraph,
    add_screenshot_placeholder,
    add_table,
    clear_document_body,
    configure_styles,
    normalize_sections,
)


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "模板-基于深度学习的废弃物数据管理系统毕业设计说明书.docx"
FALLBACK_TARGET = ROOT / "模板-基于深度学习的废弃物数据管理系统毕业设计说明书-数据集出处修正版.docx"


def add_cover(doc: Document) -> None:
    """写入封面，保留个人信息占位，方便用户按学校要求补齐。"""

    add_center(doc, "SHANDONG UNIVERSITY OF TECHNOLOGY", size=12, bold=True, font="Times New Roman")
    for _ in range(3):
        add_paragraph(doc, "", indent=False)
    add_center(doc, "毕业设计说明书", size=22, bold=True, font="黑体")
    for _ in range(3):
        add_paragraph(doc, "", indent=False)
    add_center(doc, "基于深度学习的废弃物数据管理系统", size=18, bold=True, font="黑体")
    for _ in range(4):
        add_paragraph(doc, "", indent=False)
    for item in [
        "学    院：计算机科学与技术学院",
        "专    业：计算机科学与技术",
        "学生姓名：__________",
        "学    号：__________",
        "指导教师：__________",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run(item)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体")
        run.font.size = Pt(14)
        paragraph.paragraph_format.line_spacing = 1.8
    for _ in range(2):
        add_paragraph(doc, "", indent=False)
    add_center(doc, "2026 年 5 月", size=14)
    doc.add_page_break()


def add_abstracts(doc: Document) -> None:
    """写入扩写后的中英文摘要。"""

    add_center(doc, "摘  要", size=16, bold=True, font="黑体")
    for text in [
        "随着城市化进程加快和居民消费结构持续变化，生活垃圾呈现出数量增长快、种类复杂、包装材料多样和投放场景分散等特点。垃圾分类作为城市环境治理和资源循环利用的重要环节，已经从单纯依赖宣传引导逐步转向数据化、智能化和精细化管理。传统垃圾分类方式主要依靠用户经验、静态海报和人工巡检，虽然能够满足基础宣传需求，但在面对外观相似、材质复合、污染状态不确定的废弃物时，容易出现判断标准不统一、识别效率低、可解释性不足和历史数据难以沉淀等问题。因此，研究一套能够完成图片识别、相似案例检索、分类知识查询和智能解释的废弃物数据管理系统，具有较强的现实意义。",
        "本文设计并实现了一套基于深度学习的废弃物数据管理系统。系统以 ResNet50 卷积神经网络为核心图像分类模型，围绕可回收物、有害垃圾、厨余垃圾和其他垃圾四个类别开展数据整理、模型训练与系统集成。后端采用 Flask 框架实现接口服务，负责图片上传校验、模型推理、向量检索、历史记录保存和外部 AI 接口调用；前端采用 Vue 3 构建单页应用，提供现代化工作台、图片预览、结果卡片、知识查询和历史记录等交互界面；数据库层采用 SQLite 保存本地历史数据，采用 Qdrant 向量数据库保存 2048 维图片深度特征，用于实现基于视觉相似度的案例检索。",
        "在智能交互方面，系统接入 DeepSeek 文本问答接口，用于回答用户提出的垃圾分类规则、投放建议和环保常识问题；同时接入 Kimi 多模态图片理解接口，用于分析复杂图片中的物品、材质、污染状态和可能类别。通过模型分类、向量检索和大模型解释的协同，系统不仅能够给出分类结果，还能够提供相似案例和语言化说明，从而增强用户对分类依据的理解。系统当前已根据实际演示需求移除用户登录、注册和用户管理模块，暂时隐藏知识测试入口，使毕业设计演示流程更加聚焦于识别、检索、解释和记录。",
        "实验结果表明，整理后的四分类数据集共包含 8213 张图片，其中可回收物 5586 张、有害垃圾 945 张、厨余垃圾 985 张、其他垃圾 697 张。ResNet50 模型在远端 Slurm GPU 环境中训练 10 轮，最佳验证准确率为 0.9821，测试集准确率达到 0.9830，满足课题任务书中分类准确率不低于 90% 的要求。Qdrant 集合 `waste_images` 共写入 8213 条图片向量，向量维度为 2048，距离度量为余弦相似度。功能测试覆盖图像识别、相似检索、分类知识、智能问答、图片理解和历史记录等模块，自动化测试结果为 11 个用例全部通过。系统实现结果说明，深度学习模型、向量数据库和大模型接口结合能够有效提升废弃物管理系统的智能化水平。",
    ]:
        add_paragraph(doc, text)
    add_paragraph(doc, "关键词：垃圾分类；深度学习；ResNet50；Flask；Vue；Qdrant；多模态理解", indent=False)
    doc.add_page_break()

    add_center(doc, "Abstract", size=16, bold=True, font="Times New Roman")
    for text in [
        "With the acceleration of urbanization and the continuous change of consumption patterns, municipal solid waste has become more diverse in material, appearance, packaging, and disposal scenarios. Waste classification is an important part of urban environmental governance and resource recycling. Traditional classification methods mainly depend on manual experience, static publicity materials, and offline inspection. These methods are useful for basic guidance, but they are not sufficient when waste items have similar appearances, mixed materials, or uncertain contamination states. Therefore, it is meaningful to design a waste data management system that can provide image recognition, similar case retrieval, knowledge search, and intelligent explanation.",
        "This thesis designs and implements a deep learning-based waste data management system. The system uses ResNet50 as the core image classification model and supports four classes: recyclable waste, hazardous waste, kitchen waste, and other waste. The backend is implemented with Flask and is responsible for image upload validation, model inference, vector search, history storage, and external AI API integration. The frontend is built as a Vue 3 single-page application, providing a modern workspace with image preview, result cards, knowledge search, intelligent chat, image understanding, and history views. SQLite is used for lightweight local history storage, while Qdrant is used to store 2048-dimensional image features for visual similarity search.",
        "For intelligent interaction, DeepSeek is integrated for text-based waste classification questions, and Kimi is used for multimodal image understanding. By combining model classification, vector retrieval, and natural language explanation, the system can not only output classification results but also provide similar examples and interpretable suggestions. The experimental results show that the collected four-class dataset contains 8213 images. The ResNet50 model is trained for 10 epochs on a remote Slurm GPU cluster and achieves a best validation accuracy of 0.9821 and a test accuracy of 0.9830. The Qdrant collection stores 8213 image vectors, and warm similarity-search requests take about 0.9 seconds. Functional tests cover the main modules of the system and all 11 automated test cases pass.",
    ]:
        add_paragraph(doc, text)
    add_paragraph(doc, "Keywords: Waste classification; Deep learning; ResNet50; Flask; Vue; Qdrant; Multimodal understanding", indent=False)
    doc.add_page_break()


def add_toc_placeholder(doc: Document) -> None:
    """写入目录占位，用户可在 Word 中更新为自动目录。"""

    add_center(doc, "目  录", size=16, bold=True, font="黑体")
    for line in [
        "1 绪论",
        "2 相关技术介绍",
        "3 系统需求分析",
        "4 系统总体设计",
        "5 系统详细设计与实现",
        "6 系统测试与结果分析",
        "7 总结与展望",
        "参考文献",
        "致谢",
    ]:
        add_paragraph(doc, line, indent=False)
    add_paragraph(doc, "说明：打开 Word 后，可在此处插入或更新自动目录。", indent=False)
    doc.add_page_break()


def add_full_body(doc: Document) -> None:
    """写入扩写后的论文正文，篇幅按约 16000 字毕业设计说明书组织。"""

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    for text in [
        "生活垃圾分类是城市治理体系中的基础性工作，也是资源循环利用和污染减排的重要前提。随着居民生活水平提高，日常消费品和外卖、快递、电子产品等相关废弃物数量不断增加，垃圾种类呈现出更强的复杂性。例如同一类塑料制品可能因为污染程度不同而进入不同处理流程，废旧电池、药品、灯管等有害废弃物又需要与普通生活垃圾严格区分。用户在实际投放时往往只能依靠经验进行判断，而不同地区分类规则和宣传口径也存在差异，这使得垃圾分类的准确性和持续性面临挑战。",
        "传统垃圾分类管理通常依靠人工宣传、分类桶标识和现场监督。该方式具有实施成本较低、部署简单等优点，但也存在三个明显不足。第一，识别过程依赖用户记忆，面对不常见物品时容易出错；第二，分类过程缺少数据沉淀，无法形成可复用的案例库和统计分析结果；第三，用户获得解释的渠道有限，很多时候只能知道“应该投到哪里”，却难以理解背后的分类依据。随着人工智能技术的发展，将图像识别、向量检索和智能问答引入垃圾分类场景，可以在一定程度上弥补上述不足。",
        "深度学习尤其是卷积神经网络在图像分类任务中表现突出。与传统人工特征方法相比，卷积神经网络能够从大量样本中自动学习边缘、纹理、形状和语义组合特征，对复杂背景和拍摄角度变化具有更强适应性。ResNet 系列模型通过残差连接解决深层网络训练退化问题，既能保持较强特征表达能力，又适合在中等规模数据集上进行迁移学习。因此，本文选择 ResNet50 作为废弃物图像分类模型，在四分类垃圾数据集上完成训练与部署。",
        "仅有分类结果并不能完全满足系统应用需求。在实际演示和使用中，用户还希望看到与当前图片相似的参考案例，以便判断系统结果是否可信。向量数据库为该需求提供了可行方案。系统可以将图片转化为高维特征向量，并将这些向量连同图片路径、类别等元数据写入 Qdrant。当用户上传查询图片后，系统通过相似度计算返回视觉上接近的历史图片，从而形成“分类结果 + 相似案例”的辅助决策方式。这种方式比单纯输出类别更直观，也更适合毕业设计系统展示。",
        "近年来，大语言模型和多模态模型的发展进一步提升了智能系统的解释能力。对于垃圾分类场景，用户可能会提出“过期药品属于什么垃圾”“带油污的餐盒如何处理”等自然语言问题，也可能上传一张包含多个物体的复杂图片。传统分类模型通常只能对单张主体明确的图片给出类别，而大语言模型可以根据上下文生成更细致的说明，多模态模型还可以结合图片内容进行描述和建议。因此，本文在核心识别功能之外接入 DeepSeek 和 Kimi，增强系统的智能交流和图片理解能力。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "1.2 研究目的与意义", 2)
    for text in [
        "本文的研究目的是设计并实现一套面向毕业设计演示和本地部署的废弃物数据管理系统。系统需要围绕垃圾图片构建完整处理流程，包括数据整理、模型训练、图片上传、分类识别、相似搜索、知识检索、智能问答、图片理解和历史记录管理。与只完成单一模型训练的实验不同，本文更强调模型能力与 Web 应用的结合，使深度学习结果能够通过可交互界面被用户直接使用。",
        "从应用意义看，系统可以降低用户垃圾分类学习成本。用户只需上传图片，即可获得类别、置信度和相关解释；当分类结果不够直观时，还可以查看相似案例或向智能问答模块提问。对于学校、社区或环保宣传场景，该系统也可以作为垃圾分类教学工具，用于展示人工智能在环保领域的应用方式。",
        "从工程意义看，系统覆盖了一个较完整的软件开发流程。前期需要根据任务书确定功能边界和技术路线，中期需要整理数据集、训练模型并验证效果，后期需要完成后端接口、前端页面、数据库和外部服务集成。整个过程涉及深度学习、Web 开发、数据库设计、接口调试、测试验证和文档撰写，能够体现计算机专业毕业设计的综合实践特点。",
        "从技术意义看，本文尝试将传统图像分类模型与向量数据库和大模型接口结合起来。ResNet50 负责稳定的本地分类能力，Qdrant 提供可解释性更强的相似案例检索，DeepSeek 和 Kimi 则负责知识问答和复杂图片理解。三类技术并不是相互替代关系，而是共同构成系统的智能处理链路，为废弃物数据管理提供更完整的技术方案。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "1.3 国内外研究现状", 2)
    for text in [
        "在垃圾分类图像识别方面，早期研究多采用传统机器学习方法。研究者通常先提取颜色直方图、纹理特征、形状特征或局部关键点特征，再使用支持向量机、K 近邻、随机森林等分类器进行判断。这些方法对小规模、背景单一的数据集具有一定效果，但在实际拍摄图片中，光照变化、角度变化、遮挡、污染和背景杂乱会显著影响人工特征的稳定性。",
        "深度学习出现后，图像分类研究逐渐转向端到端特征学习。AlexNet 在 ImageNet 比赛中显著提升分类准确率，VGG 通过堆叠小卷积核构建统一结构，GoogLeNet 通过 Inception 模块提高多尺度特征提取能力，ResNet 通过残差连接使更深网络训练成为可能。国内外许多垃圾分类研究基于这些模型进行迁移学习，在塑料、纸张、玻璃、金属、厨余等类别上取得较好效果。",
        "在系统应用方面，部分研究关注移动端垃圾分类小程序或 Web 管理平台，提供拍照识别、分类规则查询和积分管理等功能。也有研究将垃圾分类系统部署在智能垃圾桶或边缘设备上，通过摄像头和轻量化模型实现自动识别与投放引导。相比这些研究，本文系统更侧重毕业设计场景下的完整软件实现，强调本地可运行、功能可演示、结果可解释和文档可验证。",
        "在向量检索方面，随着非结构化数据规模扩大，Faiss、Milvus、Qdrant 等工具被广泛用于图片检索、推荐系统和语义搜索。向量检索不依赖精确关键词，而是根据特征空间距离返回相似对象，非常适合处理图片这类难以用短文本完全描述的数据。本文将 Qdrant 引入垃圾分类系统，使用户能够看到与查询图片视觉相似的历史样本，为模型结果提供直观参考。",
        "在智能交互方面，大语言模型和多模态模型正在改变传统软件系统的人机交互方式。用户不再局限于点击菜单或输入固定关键词，而可以通过自然语言提出问题。多模态模型进一步支持图片和文本联合理解，使系统能够对复杂图片进行更全面说明。本文将 DeepSeek 和 Kimi 作为系统增强能力接入，但核心识别流程仍由本地模型完成，保证系统在外部接口不可用时仍保留基本功能。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "1.4 本文主要研究内容", 2)
    for text in [
        "本文主要研究内容包括五个方面。第一，完成废弃物图片数据集整理。实验原始数据选用 Kaggle 平台 Mostafa Mohamed 发布的公开图像数据集 `Garbage Classification (12 classes)`，本文根据任务书要求将其重新映射为可回收物、有害垃圾、厨余垃圾和其他垃圾四类，并对各类别数量进行统计和检查。",
        "第二，完成 ResNet50 图像分类模型训练。本文使用 PyTorch 和 torchvision 构建训练流程，在远端 Slurm GPU 集群中运行训练任务，保存模型权重、类别映射、训练指标和分类报告，并根据测试集准确率评估模型是否满足系统要求。",
        "第三，完成系统后端设计与实现。后端基于 Flask 构建，提供图像识别、相似检索、分类知识、智能问答、图片理解和历史记录等接口，并对上传文件、模型状态、外部 API 配置和异常情况进行处理。",
        "第四，完成前端交互界面设计。系统前端采用 Vue 3 单页应用方式实现，围绕毕业设计演示需求设计首页、图像识别、相似检索、分类知识、智能问答、图片理解和历史记录视图，同时隐藏暂不使用的知识测试入口。",
        "第五，完成系统测试与结果分析。本文通过自动化测试、模型指标、Qdrant 索引验证和人工演示流程检查系统功能，并在论文中保留截图占位，方便后续补充正式运行截图。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2 相关技术介绍", 1)
    add_heading(doc, "2.1 Flask Web 框架", 2)
    for text in [
        "Flask 是一个轻量级 Python Web 框架，具有核心简洁、扩展灵活和学习成本较低等特点。与功能高度集成的大型框架相比，Flask 更适合中小型系统和原型项目。开发者可以根据实际需求选择数据库、模板引擎、认证组件和部署方式，不必被固定结构限制。",
        "本文系统采用 Flask 作为后端框架，主要原因有三点。首先，项目的核心模型训练和推理均基于 Python 生态，Flask 能够方便地与 PyTorch、Pillow、Qdrant Client 等库协同工作。其次，系统接口以 JSON 形式与前端通信，Flask 能够简洁地实现 REST 风格接口。最后，毕业设计系统需要较高可读性和可维护性，Flask 的路由组织方式便于展示业务逻辑。",
        "在具体实现中，Flask 后端负责多个关键任务：接收图片文件并校验格式和大小；调用模型服务完成分类或特征提取；调用 Qdrant 检索相似图片；读取本地知识库并返回检索结果；调用 DeepSeek 和 Kimi 外部接口；将识别历史保存到 SQLite 数据库。通过服务层拆分，系统避免将所有逻辑堆叠在路由函数中，提高了代码结构清晰度。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2.2 Vue 3 前端技术", 2)
    for text in [
        "Vue 3 是当前常用的前端渐进式框架，适合构建交互清晰、状态变化频繁的单页应用。系统前端最初可以使用原生 HTML、CSS 和 JavaScript 完成基本功能，但随着页面模块增加，纯静态写法会出现状态管理分散、界面更新重复和交互体验不一致等问题。因此，本文将前端升级为 Vue 3 驱动的单页应用。",
        "考虑到毕业设计部署环境和演示便利性，系统没有引入完整的 npm 构建流程，而是使用本地 `frontend/vendor/vue.global.prod.js` 运行时。该方案能够在保持 Vue 响应式能力的同时减少环境依赖，用户只需启动 Flask 服务即可访问完整页面。前端通过哈希路由切换不同视图，左侧导航提供首页、图像识别、相似检索、分类知识、智能问答、图片理解和历史记录入口。",
        "系统前端设计强调演示可读性和操作直观性。首页以工作台形式展示核心能力和数据概览；图像识别、相似检索和图片理解页面都支持选择图片后的即时预览；结果区域使用卡片和表格展示类别、置信度、相似度和文本解释；历史记录页面能够查看最近识别记录。整体视觉风格从传统页面调整为暗色侧边栏与浅色工作区结合的现代管理后台风格。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2.3 ResNet50 卷积神经网络", 2)
    for text in [
        "ResNet50 是 ResNet 系列中的经典模型，具有 50 层网络结构。传统深层神经网络在层数增加后可能出现训练误差反而升高的退化问题，ResNet 通过引入残差连接，使网络学习输入与输出之间的残差映射，从而缓解深层网络训练困难。残差块的核心思想是让部分信息可以跨层直接传播，既保留低层特征，又支持高层语义抽象。",
        "对于废弃物分类任务，许多类别在形状、颜色和纹理上存在明显差异。例如纸盒、塑料瓶和金属罐具有不同边缘结构和反光特征，有害垃圾中的电池和灯管也具有较独特外观。ResNet50 能够逐层提取从局部纹理到整体形状的多级特征，因此适合用于本课题的图像分类任务。",
        "本文采用迁移学习方式训练 ResNet50。迁移学习的基本思路是利用在大规模数据集上预训练得到的通用视觉特征，再针对特定任务微调模型。相比从零开始训练，迁移学习可以减少对数据规模和训练时间的依赖，尤其适合毕业设计这类数据规模有限但需要较高准确率的场景。系统将最后全连接层替换为四分类输出层，对整理后的废弃物数据集进行训练。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2.4 Qdrant 向量数据库", 2)
    for text in [
        "Qdrant 是一款面向向量检索的数据库系统，支持高维向量存储、近邻查询、Payload 过滤和集合管理。传统关系数据库擅长结构化数据查询，例如按编号、时间、类别进行筛选；而图片、文本语义和音频等非结构化数据通常需要先转化为向量，再根据向量距离判断相似程度。Qdrant 正是为这类任务提供了专门的存储和检索能力。",
        "在本文系统中，每张参考图片都会经过 ResNet50 特征提取网络，得到一个 2048 维向量。该向量保存了图片在深度特征空间中的语义表示。系统将向量写入 Qdrant，同时在 Payload 中保存图片路径、类别和文件名。用户上传查询图片时，后端生成查询向量，并在 Qdrant 中按余弦相似度返回最接近的参考图片。",
        "引入 Qdrant 的意义在于增强系统可解释性。图像分类模型输出一个类别和置信度，但用户可能并不知道模型为什么这样判断。相似检索可以展示训练数据或参考库中视觉上接近的样本，使用户通过对比图片理解结果。对于毕业设计演示而言，该模块也能体现系统不仅完成分类，而且具备数据管理和检索能力。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2.5 SQLite 数据库", 2)
    for text in [
        "SQLite 是一种轻量级嵌入式关系数据库，具有部署简单、无需独立数据库服务、读写接口稳定等特点。对于本地演示型 Web 系统，SQLite 能够以较低复杂度保存必要业务数据。本文系统当前不启用用户认证，因此数据库主要用于保存识别历史、图片路径、类别和时间等信息。",
        "使用 SQLite 的优势在于降低部署门槛。用户启动 Flask 服务后即可使用系统，不需要额外安装和配置 MySQL 或 PostgreSQL。对于毕业设计答辩环境，减少外部依赖可以降低演示失败风险。后续如果系统需要扩展为多人使用或长时间在线服务，也可以将 SQLite 替换为更适合并发访问的数据库。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2.6 DeepSeek 与 Kimi 接口", 2)
    for text in [
        "DeepSeek 是用于自然语言问答的外部大语言模型接口。本文系统将其用于垃圾分类智能交流模块，用户可以输入自然语言问题，系统把问题发送给 DeepSeek，并通过提示词限定回答范围，使其尽量围绕垃圾分类、投放建议、注意事项和环保知识进行说明。",
        "Kimi 多模态接口用于图片理解模块。与 ResNet50 的四分类输出不同，多模态模型能够结合图片和文本提示，对图片中的物体、材质、状态和可能处理方式进行更详细描述。该能力适合处理主体不明显、包含多个物品或需要解释污染状态的图片。系统将上传图片编码为 Base64 data URL 后发送给 Kimi，并将返回文本展示给用户。",
        "外部 AI 接口可以提升系统交互体验，但也带来网络依赖、密钥安全和调用成本问题。因此，本文系统将 API 密钥保存在本地 `.env` 文件中，并通过 `.gitignore` 排除，不写入代码仓库。系统也对配置缺失情况进行处理，当密钥未配置时返回明确提示，避免后端异常崩溃。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "3 系统需求分析", 1)
    add_heading(doc, "3.1 可行性分析", 2)
    for text in [
        "从技术可行性看，本文所采用的 Flask、Vue、PyTorch、Qdrant 和 SQLite 均为成熟技术，具有较完整的文档和社区资源。模型训练部分使用 PyTorch 和 torchvision，能够方便地加载 ResNet50 预训练权重并完成迁移学习。后端接口与前端页面之间通过 JSON 通信，结构清晰，便于调试。Qdrant 提供 Docker 部署方式，本地启动后即可通过 HTTP 接口访问。",
        "从经济可行性看，系统开发主要依赖开源框架和本地运行环境，不需要购买商业软件。由于用户本地没有 GPU，模型训练阶段借助远端 Slurm GPU 集群完成，训练结束后将模型权重同步回本地用于推理。外部 AI 接口主要用于演示增强，不影响本地核心识别和相似检索能力，因此系统总体成本可控。",
        "从操作可行性看，系统面向普通用户和毕业设计评审场景，操作流程应尽量简单。用户进入页面后可以通过左侧导航切换功能，在图像识别、相似检索和图片理解页面选择图片后即可查看预览和结果。系统不再要求登录注册，减少了演示步骤，使评审人员能够直接关注核心功能。",
        "从维护可行性看，系统代码按前端、后端服务、训练脚本、测试用例和文档资料分目录组织。模型权重、训练数据、Qdrant 存储和上传文件等运行态内容通过 `.gitignore` 管理，避免大文件或本地数据污染代码仓库。后续如果需要替换模型或恢复知识测试入口，也可以在现有结构上继续扩展。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "3.2 功能需求分析", 2)
    add_table(
        doc,
        "表3-1 系统功能需求表",
        ["功能模块", "输入", "输出", "说明"],
        [
            ["首页工作台", "用户访问系统", "模块入口、数据概览、功能说明", "用于毕业设计演示和功能导航"],
            ["图像识别", "废弃物图片", "类别、置信度、历史记录", "基于 ResNet50 完成四分类识别"],
            ["相似检索", "查询图片", "相似图片、相似度、参考类别", "基于 Qdrant 检索视觉相似案例"],
            ["分类知识", "关键词或类别", "分类规则、投放建议", "提供本地知识查询能力"],
            ["智能问答", "自然语言问题", "DeepSeek 文本回答", "解释垃圾分类规则和注意事项"],
            ["图片理解", "复杂图片", "Kimi 多模态分析结果", "识别图片物品、材质和处理建议"],
            ["历史记录", "识别行为数据", "历史列表和图片预览", "方便查看近期操作结果"],
        ],
    )
    for text in [
        "图像识别是系统的核心功能。用户上传图片后，系统应能够完成格式校验、图片保存、模型推理和结果展示。结果至少包括垃圾类别和置信度，并应能够写入历史记录，便于后续查询。相似检索功能需要在识别之外提供参考案例，系统应展示相似图片、相似度和参考类别，帮助用户理解检索结果。",
        "分类知识和智能问答属于解释性功能。分类知识模块基于本地规则，适合回答常见物品类别；智能问答模块基于 DeepSeek，适合回答更开放的问题。图片理解模块属于多模态增强功能，用户选择图片后应先看到预览，再提交给后端分析，避免不知道上传内容的情况。历史记录模块需要支持分页查看和图片预览，使系统具有基本数据管理能力。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "3.3 性能需求分析", 2)
    for text in [
        "系统性能需求主要体现在准确率、响应时间、稳定性和资源占用四个方面。准确率方面，任务书要求垃圾分类准确率不低于 90%，本文训练后的 ResNet50 测试准确率达到 0.9830，满足指标要求。响应时间方面，模型首次加载会产生冷启动耗时，因此系统应尽可能复用已加载模型，避免每次请求重复加载权重。",
        "稳定性方面，系统需要处理多种异常情况。例如上传文件为空、文件格式不支持、文件超过大小限制、模型权重不存在、Qdrant 服务未启动、外部 API 密钥缺失或网络异常等。对于这些情况，后端应返回明确错误信息，前端应以可读方式提示用户。资源占用方面，模型权重较大，Qdrant 向量库也需要一定磁盘空间，因此模型权重和向量库运行态不应进入普通 Git 提交。",
        "易用性也是性能体验的一部分。用户上传图片后应立即看到预览，系统返回结果后应以结构化卡片展示，而不是只输出原始 JSON。页面切换应保持顺畅，左侧导航应明确标识当前模块。考虑到答辩现场通常使用桌面浏览器演示，系统重点优化桌面端布局，同时保留基础响应式适配能力。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "3.4 非功能需求分析", 2)
    for text in [
        "安全性方面，系统需要限制上传文件类型和大小，避免异常文件影响服务运行。图片预览路由不能任意读取本机路径，只允许访问 uploads、data/raw 和 data/processed 等指定目录，减少路径穿越风险。外部 AI 密钥必须放在本地环境配置中，不写入源码和论文截图。",
        "可维护性方面，系统应保持模块边界清晰。模型推理逻辑集中在模型服务中，Qdrant 检索逻辑集中在相似检索服务中，DeepSeek 和 Kimi 调用分别封装，路由层只负责请求解析和响应组织。这样可以降低后续替换模型、调整接口或修改前端页面时的影响范围。",
        "可扩展性方面，系统当前面向四分类垃圾识别，但后续可以扩展更多细分类别。只要重新整理数据集、训练新的分类头并更新类别映射，前端和接口结构可以基本复用。相似检索模块也可以扩展 Payload 字段，例如加入图片来源、标注质量、可回收价值或处理建议等信息。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "4 系统总体设计", 1)
    add_heading(doc, "4.1 系统总体架构", 2)
    for text in [
        "系统采用浏览器/服务器架构，整体分为表现层、接口层、模型层、数据层和外部智能服务层。表现层由 Vue 前端负责，用户通过浏览器访问系统并完成图片选择、文本输入和结果查看。接口层由 Flask 提供统一 API，负责处理请求参数、调用业务服务和返回 JSON 数据。模型层由 PyTorch 和 ResNet50 组成，负责图像分类和特征提取。数据层由 SQLite 和 Qdrant 组成，分别保存结构化历史记录和高维图片向量。外部智能服务层包括 DeepSeek 和 Kimi，用于增强文本问答和图片理解能力。",
        "这种分层设计的优点是职责清晰。前端不直接接触模型文件和数据库，只通过 HTTP 接口与后端交互；后端不关心页面布局，只负责稳定提供数据；模型层和数据层可以独立测试和替换。当 Qdrant 服务暂时不可用时，图像识别功能仍然可以使用；当外部 AI 接口不可用时，本地分类和历史记录也不受影响。",
    ]:
        add_paragraph(doc, text)
    add_screenshot_placeholder(doc, "图4-1 系统首页与整体工作台界面", "http://127.0.0.1:5000/#home")
    add_table(
        doc,
        "表4-1 系统技术架构表",
        ["层次", "主要技术", "职责"],
        [
            ["表现层", "Vue 3、HTML、CSS", "展示导航、上传控件、图片预览、结果卡片和历史列表"],
            ["接口层", "Flask", "提供 REST API，进行参数校验、异常处理和响应封装"],
            ["模型层", "PyTorch、ResNet50", "完成四分类识别和 2048 维深度特征提取"],
            ["数据层", "SQLite、Qdrant", "保存历史记录和图片向量索引"],
            ["智能服务层", "DeepSeek、Kimi", "提供文字问答和多模态图片理解"],
        ],
    )

    add_heading(doc, "4.2 系统功能结构", 2)
    for text in [
        "系统功能结构围绕废弃物图片处理流程展开。用户首先可以在首页了解系统能力，然后进入图像识别模块上传图片并获得分类结果。如果用户希望验证结果或查找参考案例，可以进入相似检索模块。对于不确定的物品名称，可以使用分类知识检索；对于更复杂的问题，可以使用智能问答；对于包含多个物体或需要语义描述的图片，可以使用图片理解模块。所有识别相关行为可以在历史记录中查看。",
        "根据用户最新需求，系统已移除登录、注册和用户管理模块。这一调整符合毕业设计演示的实际情况：当前系统主要运行在本地，核心展示目标是深度学习识别和智能检索能力，而不是多用户权限管理。知识测试模块当前也暂时隐藏入口，避免答辩演示时分散重点，但后端题库接口和部分页面代码仍保留，后续可以恢复。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "4.3 后端接口设计", 2)
    add_table(
        doc,
        "表4-2 后端主要接口设计",
        ["接口路径", "方法", "功能", "说明"],
        [
            ["/api/recognize", "POST", "图像识别", "上传图片并返回类别和置信度"],
            ["/api/similar-search", "POST", "相似检索", "上传图片并返回 Qdrant 相似案例"],
            ["/api/search", "GET", "分类知识检索", "根据关键词返回本地知识内容"],
            ["/api/chat", "POST", "智能问答", "调用 DeepSeek 返回文本回答"],
            ["/api/image-understanding", "POST", "图片理解", "调用 Kimi 返回多模态分析结果"],
            ["/api/history", "GET", "历史记录", "分页返回识别历史数据"],
            ["/media/<path>", "GET", "安全图片预览", "限制读取允许目录下的图片"],
        ],
    )
    for text in [
        "接口设计遵循前后端分离思想。前端只需要根据功能选择相应接口，不需要了解模型权重路径、数据库连接和外部服务细节。后端统一进行错误处理，例如上传文件为空时返回参数错误，模型文件缺失时返回模型未就绪，外部 API 密钥未配置时返回配置提示。这种方式能够提升系统稳定性，也便于自动化测试。",
        "在图片相关接口中，系统使用 multipart/form-data 方式接收文件。后端先检查文件扩展名和大小，再保存到上传目录。对于识别接口，系统调用分类模型；对于相似检索接口，系统调用特征提取和 Qdrant 查询；对于图片理解接口，系统将图片编码为 Base64 data URL 后提交给 Kimi。不同接口共享部分上传校验逻辑，但业务处理流程各自独立。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "4.4 数据库设计", 2)
    for text in [
        "系统当前使用 SQLite 保存历史记录。由于登录注册和用户管理模块已经移除，数据库结构保持轻量化，重点记录用户上传图片后的识别结果。历史记录字段包括主键、图片路径、识别类别、置信度、创建时间等。该数据既可以用于前端历史记录展示，也可以作为后续分析用户常见查询类型的基础。",
        "SQLite 数据库文件位于本地运行目录中，属于运行态数据，不适合直接提交到 Git 仓库。系统通过初始化函数在应用启动时检查表结构，确保首次运行时能够自动创建必要表。对于毕业设计演示，这种方式降低了配置复杂度，用户不需要手动执行数据库建表脚本。",
    ]:
        add_paragraph(doc, text)
    add_table(
        doc,
        "表4-3 历史记录主要字段设计",
        ["字段", "类型", "含义"],
        [
            ["id", "INTEGER", "历史记录主键"],
            ["image_path", "TEXT", "上传图片或预览图片路径"],
            ["category", "TEXT", "模型识别出的垃圾类别"],
            ["confidence", "REAL", "模型输出置信度"],
            ["created_at", "TEXT", "记录创建时间"],
        ],
    )

    add_heading(doc, "4.5 向量索引设计", 2)
    for text in [
        "Qdrant 向量索引是相似检索模块的核心。系统通过 `scripts/build_qdrant_index.py` 扫描四分类参考图片目录，对每张图片提取 ResNet50 深度特征，并将向量写入集合 `waste_images`。向量维度为 2048，距离度量为 Cosine。Payload 中保存图片相对路径、文件名和类别，便于检索结果返回后在前端展示图片和类别。",
        "相似检索接口设置默认阈值 0.65，只有相似度超过阈值的结果才展示给用户。阈值过低会返回大量视觉上不相关的图片，影响参考价值；阈值过高则可能导致结果为空。当前阈值是在实际演示中折中选择的结果，后续可以根据更多测试样本继续调整。",
    ]:
        add_paragraph(doc, text)
    add_table(
        doc,
        "表4-4 Qdrant 向量索引配置",
        ["配置项", "取值"],
        [
            ["集合名称", "waste_images"],
            ["向量维度", "2048"],
            ["距离度量", "Cosine"],
            ["索引图片数量", "8213"],
            ["默认相似度阈值", "0.65"],
            ["本地服务地址", "http://localhost:63330"],
        ],
    )

    add_heading(doc, "4.6 安全性与异常处理设计", 2)
    for text in [
        "系统虽然主要用于本地毕业设计演示，但仍需要考虑基础安全性和异常处理。图片上传是 Web 系统中较容易出现风险的环节，因此后端需要限制允许的文件类型，只接收常见图片格式，并限制单个文件大小。对于文件名和保存路径，系统不直接信任用户输入，而是由后端生成安全路径并保存到指定上传目录，避免覆盖重要文件或访问任意路径。",
        "图片预览路由同样需要安全控制。系统提供 `/media/<path>` 用于前端展示上传图片和数据集图片，但该路由不能开放读取整个磁盘。当前实现只允许访问 `uploads`、`data/raw` 和 `data/processed` 等项目内指定目录，能够满足页面预览需求，同时降低路径穿越风险。若请求路径不在允许范围内，后端应拒绝访问。",
        "异常处理方面，系统将常见失败场景转化为可读提示。模型权重不存在时，接口返回模型未就绪；Qdrant 服务不可用时，相似检索返回检索服务异常；DeepSeek 或 Kimi 密钥缺失时，接口提示需要配置环境变量；上传文件为空或格式错误时，前端展示对应错误。通过这些处理，系统在资源不完整的情况下也能保持可解释状态，便于调试和答辩演示。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "5 系统详细设计与实现", 1)
    add_heading(doc, "5.1 数据集整理", 2)
    for text in [
        "数据集质量直接影响模型训练效果。本文实验原始数据选用 Kaggle 平台 Mostafa Mohamed 发布的公开图像数据集 `Garbage Classification (12 classes)`。该数据集面向生活垃圾图像分类任务，原始标签包含 paper、cardboard、biological、metal、plastic、green-glass、brown-glass、white-glass、clothes、shoes、batteries 和 trash 等 12 个类别。为了符合任务书四分类要求，本文将原始类别重新整理为可回收物、有害垃圾、厨余垃圾和其他垃圾四类。整理过程中需要保证目录结构清晰、文件格式可读取、每类图片数量满足训练要求，并避免无效文件进入训练流程。",
        "数据集整理后，使用检查脚本统计各类别图片数量。结果显示，可回收物 5586 张、有害垃圾 945 张、厨余垃圾 985 张、其他垃圾 697 张，共计 8213 张。可回收物数量明显多于其他类别，这是现实数据中较常见的不均衡情况。训练时需要注意类别不均衡可能导致模型更偏向样本较多类别，因此测试阶段不仅关注总体准确率，也需要查看各类别 precision、recall 和 f1-score。",
        "数据预处理包括图片读取、尺寸调整、张量转换和归一化。由于使用 ImageNet 预训练权重，输入归一化参数与 torchvision 预训练模型保持一致。训练集、验证集和测试集划分用于分别完成参数学习、训练过程选择和最终效果评估。为了保证实验可重复性，训练脚本中固定随机种子，并将类别映射保存为 JSON 文件。",
    ]:
        add_paragraph(doc, text)
    add_table(
        doc,
        "表5-1 四分类数据集统计",
        ["类别", "英文标识", "图片数量"],
        [
            ["可回收物", "recyclable", "5586"],
            ["有害垃圾", "hazardous", "945"],
            ["厨余垃圾", "kitchen", "985"],
            ["其他垃圾", "other", "697"],
            ["合计", "-", "8213"],
        ],
    )

    add_heading(doc, "5.2 模型训练实现", 2)
    for text in [
        "本地计算机没有独立 GPU，因此模型训练放在远端 Slurm 集群中完成。项目提供 `scripts/train_resnet50.slurm` 提交脚本，用于配置分区、账号、日志路径、GPU 数量、CPU 数量、内存和训练命令。训练任务实际运行在 aws 分区的 NVIDIA L40S GPU 上，使用项目本地 `.venv` 中的 Python 环境和 CUDA 12.x 兼容的 PyTorch 版本。",
        "训练脚本 `scripts/train_resnet50.py` 负责加载数据集、构建数据增强流程、初始化模型、执行训练循环和保存结果。模型主体使用 torchvision 提供的 ResNet50 预训练权重，最后分类层根据四个类别重新定义。训练过程中记录每个 epoch 的训练损失和验证准确率，并在训练结束后输出测试集准确率、分类报告、类别映射和训练指标文件。",
        "从训练日志看，模型第 1 轮验证准确率已达到 0.9586，说明迁移学习能够快速适应当前任务；随着训练进行，损失逐步下降，验证准确率在第 9 到第 10 轮达到 0.9821 左右。训练结束后测试集准确率为 0.9830，表明模型对四类废弃物图片具有较好的识别能力。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "5.3 图像识别模块实现", 2)
    for text in [
        "图像识别模块是用户最直接接触的功能。前端页面提供图片选择控件，用户选择图片后可以立即看到预览，确认后点击识别按钮。前端通过 Fetch API 将图片发送到 `/api/recognize` 接口。后端接收文件后进行格式和大小校验，避免非法文件或超大文件进入模型流程。",
        "模型服务首次调用时加载 `models/resnet50_waste.pt` 和 `models/class_map.json`。为了减少重复加载带来的耗时，服务层会缓存模型实例。图片经过 resize、center crop、归一化等预处理后输入模型，模型输出四个类别的 logits，系统通过 softmax 转换为概率分布，并选取概率最高的类别作为最终结果。",
        "识别结果返回前，后端会将图片路径、类别和置信度写入历史记录。前端收到响应后，在页面中展示类别名称、置信度和相关提示。对于模型未就绪、文件格式不支持等情况，前端展示错误信息，用户可以根据提示检查模型文件或重新选择图片。",
    ]:
        add_paragraph(doc, text)
    add_screenshot_placeholder(doc, "图5-1 图像识别页面上传图片并显示识别结果", "http://127.0.0.1:5000/#recognize")

    add_heading(doc, "5.4 相似检索模块实现", 2)
    for text in [
        "相似检索模块复用 ResNet50 的特征提取能力，但不使用最后分类层。系统取模型倒数第二层输出作为 2048 维向量，用于表示图片的深度视觉特征。相比直接比较像素，深度特征更能反映物体形状、材质和语义信息，因此适合用于相似图片查询。",
        "建索引阶段，脚本遍历 `data/raw` 下的四分类图片目录，逐张提取特征并写入 Qdrant。查询阶段，用户上传图片后，后端生成查询向量并调用 Qdrant 查询接口。考虑到 `qdrant-client` 版本差异，系统优先使用新版 `query_points()`，并保留旧版 `search()` 兜底，提高兼容性。",
        "前端展示相似结果时，不仅显示相似度数值，还显示参考图片和对应类别。这样用户可以直观看到系统认为“相似”的依据。如果查询结果为空，系统会提示没有找到超过阈值的相似案例，而不是误导用户。",
    ]:
        add_paragraph(doc, text)
    add_formula(doc, "sim(x, y) = cos(x, y) = (x · y) / (||x|| × ||y||)")
    add_screenshot_placeholder(doc, "图5-2 相似检索页面返回相似案例和相似度", "http://127.0.0.1:5000/#similar")

    add_heading(doc, "5.5 分类知识检索实现", 2)
    for text in [
        "分类知识检索模块用于解决用户知道物品名称但不确定类别的情况。系统维护一组常见废弃物知识条目，包含物品名称、所属类别和投放建议。用户输入关键词后，后端根据关键词匹配相关条目，并将结果返回给前端展示。",
        "该模块与图像识别形成互补。图像识别适合用户手中有实物图片的情况，而知识检索适合用户想快速查询文字规则的情况。比如用户输入“电池”，系统可以返回有害垃圾及其投放注意事项；用户输入“塑料瓶”，系统可以返回可回收物以及清空内容物、压扁后投放等建议。",
    ]:
        add_paragraph(doc, text)
    add_screenshot_placeholder(doc, "图5-3 分类知识检索页面，搜索电池或塑料瓶", "http://127.0.0.1:5000/#search")

    add_heading(doc, "5.6 智能问答模块实现", 2)
    for text in [
        "智能问答模块通过 DeepSeek API 实现。用户可以输入自然语言问题，前端将问题发送到 `/api/chat` 接口，后端在请求中加入系统提示词，要求模型以垃圾分类助手身份回答。提示词设计的目的是限制回答范围，减少模型生成与系统主题无关的内容。",
        "与本地知识检索相比，智能问答更适合开放问题。例如用户可以询问“过期药品属于什么垃圾”“外卖餐盒有油污还能回收吗”“碎玻璃应该如何处理”等。模型可以根据问题生成更完整的解释，说明分类原因和注意事项。系统在前端将回答以段落方式展示，便于用户阅读。",
        "由于外部 API 依赖网络和密钥配置，系统对异常情况进行了处理。当 `.env` 中没有配置 DeepSeek 密钥时，接口会返回明确提示；当网络请求失败时，系统也会返回错误说明。这样可以保证即使外部能力不可用，系统整体仍然保持可控状态。",
    ]:
        add_paragraph(doc, text)
    add_screenshot_placeholder(doc, "图5-4 智能问答页面，提问过期药品属于什么垃圾", "http://127.0.0.1:5000/#chat")

    add_heading(doc, "5.7 图片理解模块实现", 2)
    for text in [
        "图片理解模块使用 Kimi 多模态接口，主要面向复杂图片场景。传统分类模型通常假设图片主体明确、类别单一，但实际用户上传的图片可能包含多个物品、背景杂乱或物品状态不清晰。多模态模型能够根据图片内容生成文字描述，对物品名称、材质、污染状态和处理方式进行综合分析。",
        "前端在用户选择图片后立即显示预览，这是提升体验的重要细节。用户可以确认自己选择的图片是否正确，再点击分析按钮。后端接收图片后，将其编码为 Base64 data URL，并按照 OpenAI 兼容格式构造请求发送给 Kimi。返回结果通常包含图片内容描述、可能垃圾类别和投放建议。",
        "图片理解模块并不是替代 ResNet50 分类模型，而是作为解释增强能力存在。对于主体明确的图片，ResNet50 能够快速输出类别；对于复杂图片，Kimi 能够提供更丰富文本说明。二者结合可以提升系统面对真实场景的适应性。",
    ]:
        add_paragraph(doc, text)
    add_screenshot_placeholder(doc, "图5-5 图片理解页面选择图片后显示预览和 Kimi 分析结果", "http://127.0.0.1:5000/#understand")

    add_heading(doc, "5.8 历史记录模块实现", 2)
    for text in [
        "历史记录模块用于展示用户近期操作结果。每次图像识别完成后，系统将图片路径、识别类别、置信度和创建时间写入 SQLite。前端访问历史记录页面时，通过 `/api/history` 分页获取数据，并以列表或卡片形式展示。",
        "为了保证图片预览安全，后端提供 `/media/<path>` 路由，并限制只能读取指定目录下的图片。这样既能在前端展示上传图片，又避免用户通过构造路径访问系统任意文件。历史记录模块体现了系统的数据管理属性，使其不只是一次性识别工具，而是能够积累和展示识别结果的管理系统。",
    ]:
        add_paragraph(doc, text)
    add_screenshot_placeholder(doc, "图5-6 历史记录页面展示图片、类别和时间", "http://127.0.0.1:5000/#history")

    add_heading(doc, "5.9 前端界面实现", 2)
    for text in [
        "系统前端经历了从传统静态页面到 Vue 单页应用的改造。改造前，页面主要依靠原生 DOM 操作更新状态，随着功能增加，代码可读性和维护性下降。改造后，页面状态由 Vue 响应式数据驱动，不同模块通过当前视图状态切换，减少了重复代码。",
        "视觉设计方面，系统采用深色侧边导航和浅色内容区布局。深色导航增强功能入口识别度，浅色工作区保证表单、卡片和结果内容具有较好可读性。首页加入四类废弃物样例图和数据统计卡，使系统打开后能够立即传达主题和能力。按钮、输入框、结果卡片和表格均采用统一风格，提高整体一致性。",
        "根据用户反馈，系统移除了登录注册和用户管理入口，避免毕业设计演示中出现与核心主题无关的流程。知识测试模块当前也从导航中隐藏，后续如需恢复，只需重新开放入口并完善题库展示。这样的调整使系统界面更加聚焦，符合答辩展示“短时间突出核心功能”的要求。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "5.10 关键代码组织与工程实现", 2)
    for text in [
        "项目代码按照前端、后端、模型脚本、测试用例和文档资料进行组织。`waste_app` 目录保存 Flask 应用主体，其中 `routes.py` 负责接口路由，`database.py` 负责 SQLite 初始化和连接管理，`config.py` 负责读取环境变量和系统配置。`waste_app/services` 目录保存业务服务，分别封装历史记录、模型推理、相似检索、DeepSeek 调用和 Kimi 调用等逻辑。这样的结构能够避免路由文件过于庞大，也便于单独测试某一类服务。",
        "`frontend` 目录保存前端页面、样式和 Vue 运行时文件。`index.html` 提供页面挂载点和基础结构，`app.js` 保存 Vue 应用状态、视图切换、接口调用和结果渲染逻辑，`styles.css` 负责整体视觉风格和响应式布局。为了提高本地部署便利性，项目没有依赖 npm 构建流程，而是将 Vue 运行时文件放在 `frontend/vendor` 目录中，使 Flask 静态服务能够直接提供前端资源。",
        "`scripts` 目录保存数据检查、模型训练、Qdrant 建索引和论文生成相关脚本。`check_dataset.py` 用于检查整理后的数据集数量是否满足训练要求；`train_resnet50.py` 用于执行模型训练；`train_resnet50.slurm` 用于在远端 Slurm 集群提交 GPU 训练任务；`build_qdrant_index.py` 用于扫描图片并构建向量索引。论文生成脚本则用于将项目真实数据、训练结果和测试结果写入 DOCX 文档。",
        "项目测试代码位于 `tests` 目录中，主要使用 pytest 验证后端接口和异常处理。测试用例并不直接依赖真实模型权重和外部 API，而是通过未就绪提示和缺配置提示验证系统在资源缺失时的行为。这种测试方式适合毕业设计项目，因为它能够在不启动完整训练环境和不暴露真实密钥的情况下保证基础接口质量。",
        "在工程实现中，项目特别注意运行态文件和源码文件的边界。`.env`、上传图片、Qdrant 存储、训练数据和大模型权重等内容具有本地性、敏感性或体积较大，不适合进入 Git 历史；而训练指标、分类报告、论文文档和说明文件则可以作为毕业设计证据保留。通过 `.gitignore` 管理这些边界，可以减少后续提交和同步时的风险。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "6 系统测试与结果分析", 1)
    add_heading(doc, "6.1 测试环境", 2)
    add_table(
        doc,
        "表6-1 系统测试环境",
        ["项目", "配置"],
        [
            ["本地系统", "Windows 本地开发环境"],
            ["后端环境", "Python 虚拟环境、Flask"],
            ["前端环境", "Vue 3 全局运行时、浏览器"],
            ["向量数据库", "Docker Qdrant，宿主机端口 63330"],
            ["训练环境", "Slurm GPU 节点，NVIDIA L40S"],
            ["测试命令", ".venv\\Scripts\\python.exe -m pytest -q"],
        ],
    )
    for text in [
        "系统测试分为模型测试、接口功能测试、前端演示测试和相似检索性能测试。模型测试关注分类准确率和各类别指标；接口功能测试关注后端在正常输入和异常输入下的响应；前端演示测试关注页面是否能够完成完整交互流程；相似检索性能测试关注 Qdrant 索引规模、查询结果和响应时间。",
        "测试环境尽量贴近实际演示环境。后端和前端运行在本地 Windows 开发机，Qdrant 通过 Docker 容器提供服务。由于 Windows 保留了默认端口所在的 6256-6355 端口段，本地宿主机使用 63330/63340 映射容器内 6333/6334。模型训练则在远端 Slurm GPU 节点完成，训练产物同步回本地用于推理。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "6.2 功能测试", 2)
    add_table(
        doc,
        "表6-2 主要功能测试用例",
        ["编号", "测试内容", "预期结果", "测试结论"],
        [
            ["T01", "访问首页", "页面正常加载并显示核心模块", "通过"],
            ["T02", "上传合法图片识别", "返回类别和置信度", "通过"],
            ["T03", "上传非法格式文件", "返回格式错误提示", "通过"],
            ["T04", "上传超过 16MB 文件", "返回大小限制提示", "通过"],
            ["T05", "相似检索", "返回相似图片及相似度", "通过"],
            ["T06", "分类知识检索", "返回匹配知识内容", "通过"],
            ["T07", "DeepSeek 问答缺配置", "返回配置缺失提示", "通过"],
            ["T08", "Kimi 图片理解缺配置", "返回配置缺失提示", "通过"],
            ["T09", "历史记录分页", "返回分页数据", "通过"],
            ["T10", "Qdrant 客户端兼容", "新版接口可正常调用", "通过"],
        ],
    )
    for text in [
        "自动化测试使用 pytest 执行，命令为 `.venv\\Scripts\\python.exe -m pytest -q`。测试结果为 11 passed，说明当前后端基础接口和异常处理逻辑符合预期。测试覆盖健康检查、知识题库数量、文字检索、模型未就绪提示、非法格式上传、超 16MB 上传限制、DeepSeek/Kimi 缺配置错误、知识测试评分、历史记录分页删除和清空，以及 Qdrant 新版客户端相似检索调用兼容性。",
        "除了自动化测试外，还需要进行人工演示测试。原因是自动化测试主要验证接口行为，而前端页面是否美观、图片预览是否正常、结果卡片是否清晰、截图是否适合放入论文，仍需要通过浏览器实际操作确认。本文在论文中保留了多个截图占位，后续应在系统运行状态下截取正式页面。",
    ]:
        add_paragraph(doc, text)
    add_screenshot_placeholder(doc, "图6-1 pytest 自动化测试结果，显示 11 passed", "本地终端运行 .venv\\Scripts\\python.exe -m pytest -q 后截图")

    add_heading(doc, "6.3 模型测试结果", 2)
    for text in [
        "模型测试使用训练结束后保留的测试集完成。测试集共 1233 张图片，覆盖四个类别。最终测试准确率为 0.9830，说明模型整体识别能力较强。从分类报告看，可回收物类别 precision、recall 和 f1-score 均为 0.99，表现最好；有害垃圾和厨余垃圾也达到较高水平；其他垃圾 precision 为 0.93，略低于其他类别，可能与类别内部差异较大、样本数量较少有关。",
        "其他垃圾类别通常包含多种难以归入前三类的物品，其视觉特征不如可回收物或有害垃圾集中，因此模型更容易产生混淆。后续可以通过扩充其他垃圾样本、细化类别定义、增加数据增强或使用更强的模型结构进一步提升效果。尽管如此，当前宏平均 f1-score 达到 0.97，加权平均 f1-score 达到 0.98，已经满足毕业设计要求。",
    ]:
        add_paragraph(doc, text)
    add_table(
        doc,
        "表6-3 ResNet50 分类报告",
        ["类别", "Precision", "Recall", "F1-score", "Support"],
        [
            ["hazardous", "0.99", "0.98", "0.98", "142"],
            ["kitchen", "0.98", "0.96", "0.97", "134"],
            ["other", "0.93", "0.97", "0.95", "97"],
            ["recyclable", "0.99", "0.99", "0.99", "860"],
            ["accuracy", "-", "-", "0.98", "1233"],
            ["macro avg", "0.97", "0.97", "0.97", "1233"],
            ["weighted avg", "0.98", "0.98", "0.98", "1233"],
        ],
    )
    add_screenshot_placeholder(doc, "图6-2 远端训练日志或分类报告文件截图", "models/classification_report.txt 或 outputs/slurm/resnet50_33280811.out")

    add_heading(doc, "6.4 相似检索测试结果", 2)
    for text in [
        "Qdrant 向量索引构建完成后，集合 `waste_images` 共写入 8213 条图片向量。每条向量维度为 2048，距离度量为 Cosine。使用 `battery__battery1.jpg` 作为查询图片进行验证时，首次请求耗时约 14.01 秒，主要原因是模型冷启动和首次加载权重；热启动后两次请求耗时约 893.19 毫秒和 891.63 毫秒，能够满足本地演示需求。",
        "相似检索返回最高相似度为 1.0 的结果，说明查询图片在参考库中存在完全匹配或高度相似样本。实际使用中，如果用户上传的图片与参考库差异较大，返回结果可能减少或为空。该现象并不是系统错误，而是阈值过滤的正常结果。后续可以通过增加参考图片规模和降低阈值调节检索覆盖率。",
    ]:
        add_paragraph(doc, text)
    add_screenshot_placeholder(doc, "图6-3 Qdrant 集合状态或相似检索接口响应截图", "http://localhost:63330/dashboard 或相似检索页面结果")

    add_heading(doc, "6.5 智能接口测试结果", 2)
    for text in [
        "DeepSeek 和 Kimi 接口已通过本地真实密钥进行连通性验证。DeepSeek 用于文字问答，能够根据用户问题返回垃圾分类建议；Kimi 用于图片理解，能够对上传图片进行物品描述和分类建议。由于论文不能包含真实 API 密钥，测试截图应只展示页面问题、返回结果和接口成功状态，不应展示 `.env` 文件内容。",
        "外部 AI 接口的测试重点不在于模型本身的绝对正确性，而在于系统是否能够正确组装请求、处理响应和展示结果。当 API 密钥缺失时，系统应返回配置缺失提示；当接口返回异常时，系统应提示用户稍后重试。这样的异常处理能够提升系统健壮性。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "6.6 测试结果分析", 2)
    for text in [
        "综合测试结果看，系统已经实现毕业设计任务书要求的主要功能。图像识别准确率达到 98.30%，明显高于 90% 的最低要求；相似检索能够基于 Qdrant 返回参考案例；分类知识、智能问答、图片理解和历史记录模块均已具备可演示能力。前端页面经过现代化改造后，能够较好支撑答辩展示。",
        "系统仍存在一些可改进之处。首先，模型冷启动耗时较长，正式部署时应在服务启动阶段预加载模型，减少用户首次请求等待时间。其次，数据集类别分布不均衡，可回收物数量远多于其他类别，后续需要扩充少数类样本。再次，外部 AI 接口依赖网络和密钥配置，演示前需要确认 `.env` 配置正确。最后，论文截图尚未正式补入，后续需要在真实运行状态下截取清晰页面。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "6.7 部署运行与演示检查", 2)
    for text in [
        "系统本地运行时需要先确认 Python 虚拟环境、模型文件、Qdrant 容器和环境变量配置。后端启动前，应保证 `.venv` 中已经安装 requirements 所需依赖，`models` 目录下存在 `resnet50_waste.pt` 和 `class_map.json`，并且 Qdrant 容器已经监听 `http://localhost:63330`。如果需要使用 DeepSeek 和 Kimi 功能，还需要在 `.env` 中配置对应 API Key。由于 `.env` 不进入 Git 仓库，换机运行时必须重新配置。",
        "答辩演示时建议按照固定顺序操作。首先打开首页展示系统主题、数据规模和模块入口；其次进入图像识别页面上传一张主体清晰的废弃物图片，展示类别和置信度；然后进入相似检索页面上传同类图片，展示 Qdrant 返回的参考案例；接着演示分类知识和智能问答，说明系统不仅能识别图片，也能解释规则；最后演示图片理解和历史记录，展示多模态分析和数据管理能力。",
        "为了保证演示稳定，建议提前完成模型预热。可以先在浏览器中执行一次图像识别请求，使后端加载模型权重，后续请求响应会更快。相似检索同样可以提前用一张已知图片测试，确认 Qdrant 容器、集合名称和向量维度均正常。如果现场网络不稳定，DeepSeek 和 Kimi 可能受到影响，因此论文和答辩中应说明外部 AI 是增强能力，本地 ResNet50 识别和 Qdrant 检索才是系统核心功能。",
        "论文截图应选择清晰、信息完整且不暴露敏感配置的画面。首页截图需要体现系统名称和模块入口；识别截图需要包含上传图片、类别和置信度；相似检索截图需要包含相似图片和相似度；智能问答截图需要展示问题和回答；图片理解截图需要展示图片预览和分析文本；测试截图需要展示 11 passed；训练截图应展示测试准确率或分类报告。截图插入 Word 后应统一宽度，图题位于图片下方，并保持正文排版整齐。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "7 总结与展望", 1)
    add_heading(doc, "7.1 总结", 2)
    for text in [
        "本文围绕“基于深度学习的废弃物数据管理系统”完成了从需求分析、技术选型、模型训练、系统设计、功能实现到测试验证的完整工作。系统以 ResNet50 为核心分类模型，结合 Flask 后端、Vue 前端、SQLite 数据库和 Qdrant 向量数据库，实现了图像识别、相似检索、分类知识、智能问答、图片理解和历史记录等功能。",
        "在模型训练方面，本文整理了 8213 张四分类废弃物图片，并在远端 Slurm GPU 环境中完成 ResNet50 迁移训练。最终模型测试准确率达到 0.9830，满足任务书指标。在系统实现方面，本文将模型推理能力封装为 Web 接口，并通过现代化前端页面提供可交互演示。在数据管理方面，系统使用 SQLite 保存历史记录，使用 Qdrant 保存图片向量并实现相似检索。在智能交互方面，系统接入 DeepSeek 和 Kimi，提升了分类解释和复杂图片理解能力。",
        "本文工作说明，深度学习模型并不应只停留在离线训练结果上，而应与系统工程结合，形成可使用、可展示、可测试的软件应用。通过本课题实践，可以更加直观地理解人工智能技术在环保信息化场景中的应用方式，也能够锻炼模型训练、接口开发、前端设计、数据库管理和文档撰写等综合能力。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "7.2 展望", 2)
    for text in [
        "后续工作可以从数据、模型、系统和应用四个方向继续优化。数据方面，可以继续采集真实投放场景图片，特别是增加其他垃圾、有害垃圾和混合物品图片，缓解类别不均衡问题。还可以建立更细粒度标签体系，在四大类基础上增加塑料、纸张、金属、电池、药品、厨余残渣等子类别。",
        "模型方面，可以尝试 EfficientNet、ConvNeXt、Vision Transformer 或轻量化 MobileNet 等结构，对比准确率、推理速度和模型大小。对于本地部署场景，可以考虑模型量化和 ONNX 导出，减少推理资源占用。对于误分类样本，可以建立错误分析流程，定期将高价值样本加入训练集继续微调。",
        "系统方面，可以进一步完善 Qdrant 数据管理工具，支持索引增量更新、重复图片检查、检索结果反馈和向量库可视化。前端可以恢复并完善知识测试模块，形成“识别—学习—测试—复习”的闭环。部署方面，可以引入模型预热、缓存、异步任务和日志监控，提高系统长期运行稳定性。",
        "应用方面，系统可以扩展到校园、社区或企业环保宣传场景。通过接入摄像头、移动端页面或智能垃圾桶硬件，系统能够提供更贴近真实使用的分类辅助能力。同时，历史数据也可以用于分析用户常见困惑物品，为后续垃圾分类宣传和规则优化提供参考。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "7.3 论文材料组织说明", 2)
    for text in [
        "为了使论文材料与系统实现相互对应，本文在正文中尽量使用项目已经验证过的数据和结果。数据集数量来自检查脚本输出，模型准确率和分类报告来自训练完成后的 `models` 目录，Qdrant 索引规模和查询耗时来自本地构建与接口验证，自动化测试结果来自 pytest 命令。这些材料可以支撑系统测试章节，避免论文只停留在功能描述层面。",
        "截图材料需要在系统最终运行状态下补充。由于截图涉及浏览器页面、终端输出和 Qdrant Dashboard，不适合在生成论文时直接固定到文档中，因此本文先保留截图占位。后续补图时，应确保截图内容与正文描述一致，例如图像识别页面应显示真实上传图片和分类置信度，图片理解页面应显示预览图和 Kimi 返回结果，测试截图应显示 11 passed。",
        "正式提交前还需要统一检查论文格式。重点包括封面个人信息是否补齐、目录是否更新、章节标题是否连续、表格是否跨页错乱、截图是否清晰、图题表题是否规范、参考文献格式是否统一。由于本文基于既有 DOCX 模板生成，页面基础格式已经继承模板，但插入截图后仍可能改变分页，因此最终版必须在 Word 中人工复核。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "参考文献", 1)
    for item in [
        "[1] He K, Zhang X, Ren S, et al. Deep Residual Learning for Image Recognition[C]. IEEE Conference on Computer Vision and Pattern Recognition, 2016.",
        "[2] Paszke A, Gross S, Massa F, et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library[C]. NeurIPS, 2019.",
        "[3] Pallets Projects. Flask Documentation[EB/OL]. https://flask.palletsprojects.com/.",
        "[4] Qdrant. Qdrant Vector Database Documentation[EB/OL]. https://qdrant.tech/documentation/.",
        "[5] Vue.js Team. Vue 3 Documentation[EB/OL]. https://vuejs.org/.",
        "[6] DeepSeek. DeepSeek API Documentation[EB/OL]. https://api-docs.deepseek.com/.",
        "[7] Moonshot AI. Kimi API Documentation[EB/OL]. https://platform.moonshot.cn/docs/.",
        "[8] Mohamed M. Garbage Classification (12 classes)[DB/OL]. Kaggle, 2021. https://www.kaggle.com/datasets/mostafaabla/garbage-classification.",
        "[9] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016.",
        "[10] Krizhevsky A, Sutskever I, Hinton G E. ImageNet Classification with Deep Convolutional Neural Networks[J]. Communications of the ACM, 2017.",
        "[11] Johnson J, Douze M, Jégou H. Billion-scale similarity search with GPUs[J]. IEEE Transactions on Big Data, 2019.",
    ]:
        add_paragraph(doc, item, indent=False)

    add_heading(doc, "致谢", 1)
    for text in [
        "在本课题完成过程中，感谢指导教师在选题确定、技术路线、系统实现和论文撰写等方面给予的指导与帮助。老师的建议使我能够从单纯模型训练扩展到完整系统实现，也让我更加重视需求分析、测试验证和文档规范。",
        "感谢学院提供的学习环境和实践条件，感谢课程学习中积累的编程、数据库、软件工程和人工智能相关知识，为本课题顺利完成打下基础。感谢开源社区提供的 Flask、Vue、PyTorch、Qdrant 等优秀工具，使个人开发者也能够构建较完整的智能 Web 系统。",
        "通过本次毕业设计，我进一步理解了深度学习技术从实验结果到软件系统落地之间的差距，也提升了问题分析、代码实现、环境配置、远端训练、接口调试和论文写作能力。今后将继续保持对人工智能应用和软件工程实践的学习，不断提高系统设计与工程实现水平。",
    ]:
        add_paragraph(doc, text)


def main() -> None:
    if not TARGET.exists():
        raise FileNotFoundError(f"未找到目标模板文件：{TARGET}")

    doc = Document(TARGET)
    clear_document_body(doc)
    configure_styles(doc)
    normalize_sections(doc)
    add_cover(doc)
    add_abstracts(doc)
    add_toc_placeholder(doc)
    add_full_body(doc)
    try:
        doc.save(TARGET)
        print(f"已写入扩写论文：{TARGET}")
    except PermissionError:
        # 如果用户正在用 Word/WPS 打开目标文件，Windows 会锁定 DOCX。
        # 此时保存到修正版副本，避免中断论文内容更新；用户关闭原文件后可再次运行本脚本覆盖目标文件。
        doc.save(FALLBACK_TARGET)
        print(f"目标文件被占用，已写入修正版副本：{FALLBACK_TARGET}")


if __name__ == "__main__":
    main()
