"""根据既有毕业设计说明书模板生成本项目论文 DOCX。

本脚本不会覆盖原模板文件，而是读取模板中的页面设置和基础样式，
生成《基于深度学习的废弃物数据管理系统》毕业设计说明书。正文中
需要后期补充系统截图的位置会以“截图占位”段落标出，便于用户在
Word 中手动替换为真实截图。
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "21110543069王逸飞基于遗传算法的旅游路线规划系统.docx"
OUTPUT = ROOT / "基于深度学习的废弃物数据管理系统毕业设计说明书.docx"


def clear_document_body(doc: Document) -> None:
    """清空模板正文内容，同时保留最后一个 section 的页面属性。

    Word 文档正文的最后一个元素通常是 `sectPr`，它保存页边距、纸张大小、
    页眉页脚等页面设置。清空正文时保留该元素，可以让新论文继续沿用模板
    的版式基础，而不会破坏文档的页面结构。
    """

    body = doc._body._element
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = deepcopy(child)
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def set_east_asia_font(run, font_name: str = "宋体", size: Pt | None = None) -> None:
    """统一设置中英文字体，避免中文在不同 Word 环境中回退异常。"""

    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = size


def configure_styles(doc: Document) -> None:
    """配置论文常用样式，使新内容贴近本科毕业设计说明书排版习惯。"""

    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, *, indent: bool = True):
    """添加正文段落，并按论文正文要求设置首行缩进和行距。"""

    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        set_east_asia_font(run)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if indent and style is None:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    else:
        paragraph.paragraph_format.first_line_indent = None
    return paragraph


def add_center(doc: Document, text: str, *, size: int = 12, bold: bool = False, font: str = "宋体"):
    """添加居中段落，主要用于封面、摘要标题、图题和表题。"""

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_east_asia_font(run, font, Pt(size))
    run.bold = bold
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.5
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    """添加章节标题，并复用 Word 的 Heading 样式，便于后续自动生成目录。"""

    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run(text)
    set_east_asia_font(run, "黑体", Pt(16 if level == 1 else 14 if level == 2 else 12))
    run.bold = True
    return paragraph


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    """添加论文表格，表题置于表格上方，内容使用紧凑但可读的排版。"""

    add_center(doc, title, size=10, bold=False)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_east_asia_font(run, "宋体", Pt(10.5))
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_east_asia_font(run, "宋体", Pt(10.5))
    add_paragraph(doc, "", indent=False)


def add_screenshot_placeholder(doc: Document, caption: str, page_hint: str) -> None:
    """添加截图占位说明，后续用户可直接在该位置插入真实系统截图。"""

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"【截图占位：{caption}。后期请截取：{page_hint}】")
    set_east_asia_font(run, "黑体", Pt(11))
    run.font.color.rgb = RGBColor(192, 0, 0)


def add_formula(doc: Document, text: str) -> None:
    """以居中方式添加简化公式说明，适合论文中展示核心计算关系。"""

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run(text)
    set_east_asia_font(run, "Times New Roman", Pt(11))


def add_cover(doc: Document) -> None:
    """生成封面信息，未知个人信息保留占位，方便用户后续按学校要求补齐。"""

    add_center(doc, "SHANDONG UNIVERSITY OF TECHNOLOGY", size=12, bold=True, font="Times New Roman")
    for _ in range(3):
        add_paragraph(doc, "", indent=False)
    add_center(doc, "毕业设计说明书", size=22, bold=True, font="黑体")
    for _ in range(3):
        add_paragraph(doc, "", indent=False)
    add_center(doc, "基于深度学习的废弃物数据管理系统", size=18, bold=True, font="黑体")
    for _ in range(4):
        add_paragraph(doc, "", indent=False)
    cover_items = [
        "学    院：计算机科学与技术学院",
        "专    业：计算机科学与技术",
        "学生姓名：__________",
        "学    号：__________",
        "指导教师：__________",
    ]
    for item in cover_items:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(item)
        set_east_asia_font(run, "宋体", Pt(14))
        paragraph.paragraph_format.line_spacing = 1.8
    for _ in range(2):
        add_paragraph(doc, "", indent=False)
    add_center(doc, "2026 年 5 月", size=14)
    doc.add_page_break()


def add_abstracts(doc: Document) -> None:
    """生成中英文摘要和关键词。"""

    add_center(doc, "摘  要", size=16, bold=True, font="黑体")
    add_paragraph(
        doc,
        "随着城市生活垃圾产生量持续增长，垃圾分类工作对识别准确率、处理效率和知识普及提出了更高要求。"
        "传统垃圾分类方式主要依赖人工经验判断，存在识别标准不统一、用户学习成本较高、历史数据难以积累等问题。"
        "针对上述问题，本文设计并实现了一套基于深度学习的废弃物数据管理系统，将图像分类、向量检索、知识检索、"
        "智能问答和多模态图片理解等功能整合到统一的 Web 工作台中，为用户提供从图片上传、分类识别到相似案例查询和分类解释的完整流程。",
    )
    add_paragraph(
        doc,
        "系统后端采用 Flask 框架实现接口服务，前端采用 Vue 3 构建单页交互界面。图像识别部分以 ResNet50 卷积神经网络为核心模型，"
        "在整理后的四分类垃圾图片数据集上完成训练，覆盖可回收物、有害垃圾、厨余垃圾和其他垃圾四类场景。相似检索模块复用 ResNet50 "
        "的深度特征表示，将 2048 维图像向量写入 Qdrant 向量数据库，并基于余弦相似度返回相似案例。系统同时引入 DeepSeek 文本问答接口"
        "和 Kimi 多模态图片理解接口，增强系统对垃圾分类知识解释、复杂图片分析和用户自然语言咨询的支持能力。",
    )
    add_paragraph(
        doc,
        "实验结果表明，ResNet50 模型在测试集上的准确率达到 0.9830，能够满足课题任务书中分类准确率不低于 90% 的要求。"
        "Qdrant 集合共写入 8213 条图片向量，热启动后的相似检索响应时间约为 0.9 秒。系统功能测试覆盖图像识别、相似检索、"
        "文字检索、智能问答、图片理解和历史记录等模块，测试结果表明系统能够稳定完成主要业务流程。本文的研究与实现说明，"
        "深度学习模型与向量数据库结合能够有效提升废弃物数据管理系统的智能化水平，具有一定工程实践价值和推广意义。",
    )
    add_paragraph(doc, "关键词：垃圾分类；深度学习；ResNet50；Flask；Vue；Qdrant；多模态理解", indent=False)
    doc.add_page_break()

    add_center(doc, "Abstract", size=16, bold=True, font="Times New Roman")
    add_paragraph(
        doc,
        "With the continuous growth of municipal solid waste, waste classification systems require higher recognition accuracy, better processing efficiency, "
        "and clearer knowledge guidance. Traditional classification mainly depends on manual experience, which often leads to inconsistent criteria, "
        "high learning cost for users, and weak accumulation of historical data. To address these problems, this thesis designs and implements a deep "
        "learning-based waste data management system. The system integrates image classification, vector similarity search, knowledge retrieval, "
        "intelligent question answering, and multimodal image understanding into a unified Web workspace.",
    )
    add_paragraph(
        doc,
        "The backend is implemented with Flask, while the frontend is built as a Vue 3 single-page application. ResNet50 is used as the core image "
        "classification model and trained on a four-class waste image dataset, including recyclable waste, hazardous waste, kitchen waste, and other waste. "
        "For similarity search, the system extracts 2048-dimensional deep features from images and stores them in Qdrant, where cosine similarity is used "
        "to retrieve visually similar cases. DeepSeek is integrated for text-based waste classification consultation, and Kimi is used for multimodal "
        "image understanding.",
    )
    add_paragraph(
        doc,
        "The experimental results show that the trained ResNet50 model achieves a test accuracy of 0.9830, satisfying the requirement of no less than 90% "
        "classification accuracy. The Qdrant collection stores 8213 image vectors, and warm similarity-search requests take about 0.9 seconds. Functional "
        "tests cover image recognition, similarity search, text retrieval, intelligent chat, image understanding, and history management. The results "
        "demonstrate that the proposed system can complete the main workflows reliably and provide a practical solution for intelligent waste data management.",
    )
    add_paragraph(
        doc,
        "Keywords: Waste classification; Deep learning; ResNet50; Flask; Vue; Qdrant; Multimodal understanding",
        indent=False,
    )
    doc.add_page_break()


def add_toc_placeholder(doc: Document) -> None:
    """添加目录占位，用户可在 Word 中更新为自动目录。"""

    add_center(doc, "目  录", size=16, bold=True, font="黑体")
    for line in [
        "1 绪论",
        "2 相关技术介绍",
        "3 系统需求分析",
        "4 系统总体设计",
        "5 系统详细设计与实现",
        "6 系统测试与结果分析",
        "7 总结与展望",
        "参考文献",
        "致谢",
    ]:
        add_paragraph(doc, line, indent=False)
    add_paragraph(doc, "说明：打开 Word 后，可在此处插入或更新自动目录。", indent=False)
    doc.add_page_break()


def add_body(doc: Document) -> None:
    """写入论文主体章节。"""

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_paragraph(
        doc,
        "生活垃圾分类是城市精细化治理和绿色低碳发展的重要组成部分。随着居民消费结构变化和包装材料种类增多，日常废弃物呈现出来源复杂、"
        "外观相似、材质多样等特点。用户在投放垃圾时往往需要在短时间内判断物品类别，但不同地区分类规则、物品污染状态和回收条件存在差异，"
        "仅依靠人工记忆和静态宣传材料难以满足高频、准确、可解释的分类需求。",
    )
    add_paragraph(
        doc,
        "近年来，深度学习在图像识别领域快速发展，卷积神经网络能够从图像中自动提取纹理、形状和局部结构等特征，在垃圾分类、工业质检、"
        "医疗影像和智能安防等场景中得到广泛应用。与此同时，向量数据库能够对高维特征进行快速索引和相似检索，大语言模型和多模态模型则提升了"
        "系统对自然语言问题和复杂图片语义的解释能力。因此，将深度学习、向量检索和智能交互技术结合起来，构建废弃物数据管理系统，具有较强的应用价值。",
    )
    add_heading(doc, "1.2 研究意义", 2)
    add_paragraph(
        doc,
        "本文所设计的系统不仅关注单张图片的分类结果，还将识别结果、相似案例、知识说明和历史记录进行统一管理。对普通用户而言，系统可以降低垃圾分类"
        "学习成本，帮助用户理解分类依据；对学校或社区等管理场景而言，系统可以积累识别数据和查询记录，为后续宣传、统计和模型优化提供数据基础；"
        "对毕业设计工程实践而言，该系统覆盖模型训练、后端接口、前端交互、数据库管理和外部 AI 接口调用等多个环节，能够体现软件工程与人工智能技术的综合应用能力。",
    )
    add_heading(doc, "1.3 国内外研究现状", 2)
    add_paragraph(
        doc,
        "在图像分类研究方面，早期方法通常依赖 SIFT、HOG、颜色直方图等人工特征，并配合支持向量机、随机森林等传统分类器完成识别。该类方法对光照、"
        "背景和拍摄角度较为敏感，特征表达能力有限。深度学习兴起后，AlexNet、VGG、GoogLeNet 和 ResNet 等模型显著提升了图像分类性能。"
        "其中 ResNet 通过残差连接缓解深层网络训练退化问题，在中小规模迁移学习任务中具有稳定表现。",
    )
    add_paragraph(
        doc,
        "在向量检索方面，Faiss、Milvus、Qdrant 等工具能够对文本、图片和音频等非结构化数据的向量表示进行高效检索。相比传统关键词检索，"
        "向量检索更适合处理“视觉上相似但文件名或文字描述不同”的图片数据。当前许多智能应用也开始将图像特征检索与大语言模型解释结合起来，"
        "使系统不仅能给出结果，还能说明原因并提供进一步操作建议。",
    )
    add_heading(doc, "1.4 本文主要研究内容", 2)
    add_paragraph(
        doc,
        "本文围绕“基于深度学习的废弃物数据管理系统”展开，主要研究内容包括：第一，整理四分类废弃物图片数据集，并完成 ResNet50 分类模型训练；"
        "第二，设计 Flask 后端接口，封装图片上传、模型推理、相似检索、文字检索、智能问答、图片理解和历史记录等功能；第三，使用 Qdrant 构建图片向量库，"
        "实现基于深度特征的相似案例检索；第四，使用 Vue 3 实现现代化前端页面，提升毕业设计演示效果；第五，通过功能测试和模型测试验证系统可用性。",
    )

    add_heading(doc, "2 相关技术介绍", 1)
    add_heading(doc, "2.1 Flask 框架", 2)
    add_paragraph(
        doc,
        "Flask 是一个轻量级 Python Web 框架，具有结构灵活、扩展方便和上手成本低等特点。本文使用 Flask 实现 REST 风格接口，负责接收前端上传文件、"
        "调用模型服务、访问 SQLite 数据库、请求外部 AI API，并将结果以 JSON 格式返回给前端。由于毕业设计系统规模适中，Flask 能够在保持代码简洁的同时满足业务扩展需要。",
    )
    add_heading(doc, "2.2 Vue 3 前端技术", 2)
    add_paragraph(
        doc,
        "Vue 3 是一种渐进式 JavaScript 前端框架，适合构建响应式单页应用。系统前端采用 Vue 全局运行时方式实现，无需额外引入复杂构建流程，"
        "便于在毕业设计演示环境中直接部署。前端通过组件化状态管理组织不同业务视图，并使用 Fetch API 与 Flask 后端交互。",
    )
    add_heading(doc, "2.3 ResNet50 深度学习模型", 2)
    add_paragraph(
        doc,
        "ResNet50 是一种具有 50 层网络结构的残差卷积神经网络。残差结构通过恒等映射将浅层特征直接传递到深层，缓解了深层网络训练中的梯度消失和性能退化问题。"
        "本文在预训练 ResNet50 基础上进行迁移学习，将最后分类层替换为四分类输出层，从而适配废弃物分类任务。",
    )
    add_heading(doc, "2.4 Qdrant 向量数据库", 2)
    add_paragraph(
        doc,
        "Qdrant 是面向向量相似检索的数据库系统，支持向量存储、Payload 元数据管理和基于距离度量的近邻查询。本文将 ResNet50 提取的 2048 维图片特征写入 Qdrant，"
        "并保存图片路径、类别和文件名等元数据。用户上传查询图片后，系统提取其深度特征并在 Qdrant 中检索相似图片，从而提供可参考的历史案例。",
    )
    add_heading(doc, "2.5 大语言模型与多模态理解", 2)
    add_paragraph(
        doc,
        "大语言模型能够根据自然语言问题生成解释性回答，多模态模型则可以同时理解图片和文本。本文使用 DeepSeek 接入文字问答能力，用于回答垃圾分类规则、投放建议等问题；"
        "使用 Kimi 多模态接口对上传图片进行语义分析，补充模型分类结果之外的物品识别、材质判断和处置建议。外部 API 密钥统一通过本地环境变量配置，不写入代码仓库。",
    )

    add_heading(doc, "3 系统需求分析", 1)
    add_heading(doc, "3.1 可行性分析", 2)
    add_paragraph(
        doc,
        "从技术可行性看，系统采用的 Flask、Vue、PyTorch、Qdrant 和 SQLite 均为成熟工具，能够满足毕业设计开发和本地演示需要。ResNet50 具有成熟的预训练权重，"
        "在中等规模图片数据集上迁移训练成本可控。Qdrant 提供 Docker 部署方式，便于在本地快速构建向量检索服务。",
    )
    add_paragraph(
        doc,
        "从经济可行性看，系统主要依赖开源框架和本地运行环境，训练阶段借助远端 Slurm GPU 集群完成，避免本地无 GPU 环境下长时间训练的问题。外部 AI API 仅用于演示和增强解释，"
        "核心图像识别和相似检索能力可在本地独立运行。",
    )
    add_heading(doc, "3.2 功能需求", 2)
    add_table(
        doc,
        "表3-1 系统功能需求表",
        ["功能模块", "输入", "输出", "说明"],
        [
            ["图像识别", "用户上传废弃物图片", "分类结果、置信度、保存记录", "基于 ResNet50 完成四分类识别"],
            ["相似检索", "用户上传查询图片", "相似图片、相似度、参考类别", "基于 Qdrant 检索视觉相似案例"],
            ["分类知识", "关键词或类别", "分类规则和投放建议", "用于查询常见垃圾分类知识"],
            ["智能问答", "自然语言问题", "DeepSeek 生成回答", "用于解释分类依据和处理建议"],
            ["图片理解", "用户上传图片", "Kimi 多模态分析文本", "识别复杂图片中的物品和场景"],
            ["历史记录", "识别与检索行为", "历史列表和图片预览", "便于查看近期操作记录"],
        ],
    )
    add_heading(doc, "3.3 性能需求", 2)
    add_paragraph(
        doc,
        "系统性能需求主要包括识别准确率、响应时间、稳定性和易用性四个方面。模型分类准确率应不低于任务书要求的 90%；常规上传图片应能在可接受时间内返回结果；"
        "当模型文件、Qdrant 服务或外部 API 缺失时，系统应给出明确错误提示而不是直接崩溃；前端页面应具有清晰导航和直观操作流程，满足毕业设计现场演示需要。",
    )
    add_heading(doc, "3.4 非功能需求", 2)
    add_paragraph(
        doc,
        "系统需要具备一定安全性和可维护性。上传文件应限制格式和大小，避免异常文件影响服务运行；外部 API 密钥应通过 `.env` 文件配置，并由 `.gitignore` 排除；"
        "模型权重文件体积较大，不直接进入 Git 历史，而通过 Release 或外部制品通道同步；后端服务模块划分清晰，便于后续替换模型或扩展功能。",
    )

    add_heading(doc, "4 系统总体设计", 1)
    add_heading(doc, "4.1 系统架构设计", 2)
    add_paragraph(
        doc,
        "系统采用浏览器/服务器架构。用户通过浏览器访问 Vue 前端页面，前端负责页面渲染、文件选择、图片预览、表单提交和结果展示；Flask 后端负责业务逻辑处理；"
        "PyTorch 模型服务完成图片分类和特征提取；SQLite 保存历史记录；Qdrant 保存图片向量并完成相似检索；DeepSeek 和 Kimi 作为外部 AI 服务提供智能解释能力。",
    )
    add_screenshot_placeholder(doc, "图4-1 系统首页与整体工作台界面", "http://127.0.0.1:5000/#home")
    add_table(
        doc,
        "表4-1 系统技术架构表",
        ["层次", "主要技术", "职责"],
        [
            ["表现层", "Vue 3、HTML、CSS", "展示导航、上传控件、结果卡片和历史列表"],
            ["接口层", "Flask", "提供 REST API，处理请求校验和响应封装"],
            ["模型层", "PyTorch、ResNet50", "完成图像分类和 2048 维特征提取"],
            ["数据层", "SQLite、Qdrant", "保存历史记录和图片向量索引"],
            ["智能服务层", "DeepSeek、Kimi", "提供文字问答和多模态图片理解"],
        ],
    )
    add_heading(doc, "4.2 功能结构设计", 2)
    add_paragraph(
        doc,
        "系统功能围绕“识别、检索、解释、记录”四个核心环节组织。图像识别模块负责生成基础分类结果；相似检索模块提供参考案例；分类知识和智能问答模块负责解释规则；"
        "图片理解模块面向复杂图片提供补充判断；历史记录模块保存用户操作结果，便于后续查看和展示。",
    )
    add_heading(doc, "4.3 数据库设计", 2)
    add_paragraph(
        doc,
        "系统使用 SQLite 保存本地业务数据。当前运行版本已移除用户登录、注册和用户管理模块，因此不再创建用户表，主要保留识别历史相关数据。历史记录包含图片路径、"
        "识别类别、置信度、创建时间等字段，用于前端历史记录页面展示。由于系统面向本地毕业设计演示，SQLite 能够满足轻量级存储需求。",
    )
    add_table(
        doc,
        "表4-2 历史记录主要字段设计",
        ["字段", "类型", "含义"],
        [
            ["id", "INTEGER", "历史记录主键"],
            ["image_path", "TEXT", "上传图片或预览图片路径"],
            ["category", "TEXT", "模型识别出的垃圾类别"],
            ["confidence", "REAL", "模型输出置信度"],
            ["created_at", "TEXT", "记录创建时间"],
        ],
    )
    add_heading(doc, "4.4 向量索引设计", 2)
    add_paragraph(
        doc,
        "Qdrant 集合名称设置为 `waste_images`。系统为每张参考图片提取 ResNet50 倒数第二层输出的 2048 维特征向量，并使用余弦相似度作为距离度量。"
        "Payload 中保存类别、文件名和图片路径等信息，接口层默认过滤相似度低于 0.65 的结果，避免展示参考价值较低的图片。",
    )
    add_table(
        doc,
        "表4-3 Qdrant 向量索引配置",
        ["配置项", "取值"],
        [
            ["集合名称", "waste_images"],
            ["向量维度", "2048"],
            ["距离度量", "Cosine"],
            ["索引图片数量", "8213"],
            ["默认相似度阈值", "0.65"],
            ["本地服务地址", "http://localhost:63330"],
        ],
    )

    add_heading(doc, "5 系统详细设计与实现", 1)
    add_heading(doc, "5.1 数据集整理与模型训练", 2)
    add_paragraph(
        doc,
        "实验原始数据选用 Kaggle 平台 Mostafa Mohamed 发布的公开图像数据集 `Garbage Classification (12 classes)`，原始数据包含多个细分类别。根据课题任务书要求，本文将数据整理为可回收物、有害垃圾、厨余垃圾和其他垃圾四类。"
        "整理后的数据集共 8213 张图片，其中可回收物 5586 张、有害垃圾 945 张、厨余垃圾 985 张、其他垃圾 697 张。",
    )
    add_table(
        doc,
        "表5-1 四分类数据集统计",
        ["类别", "英文标识", "图片数量"],
        [
            ["可回收物", "recyclable", "5586"],
            ["有害垃圾", "hazardous", "945"],
            ["厨余垃圾", "kitchen", "985"],
            ["其他垃圾", "other", "697"],
            ["合计", "-", "8213"],
        ],
    )
    add_paragraph(
        doc,
        "模型训练在远端 Slurm 集群中完成。由于本地计算机没有独立 GPU，训练任务提交到具备 NVIDIA L40S 显卡的 GPU 节点运行。训练脚本采用 PyTorch 与 torchvision，"
        "加载 ImageNet 预训练 ResNet50 权重，并将分类头改为四分类输出。训练共进行 10 个 epoch，保存最终模型权重、类别映射、训练指标和分类报告。",
    )
    add_heading(doc, "5.2 图像识别模块实现", 2)
    add_paragraph(
        doc,
        "图像识别接口接收前端上传的图片文件，首先检查文件类型和大小，然后将图片保存到上传目录。模型服务加载 `models/resnet50_waste.pt` 和 `models/class_map.json`，"
        "对图片进行尺寸调整、归一化和张量转换，随后调用 ResNet50 输出四类概率分布。后端选取概率最高的类别作为识别结果，并将类别、置信度和图片路径写入历史记录。",
    )
    add_screenshot_placeholder(doc, "图5-1 图像识别页面上传图片并显示识别结果", "http://127.0.0.1:5000/#recognize")
    add_heading(doc, "5.3 相似检索模块实现", 2)
    add_paragraph(
        doc,
        "相似检索模块与图像识别模块共享图片预处理和特征提取流程，但不使用最终分类层，而是取 ResNet50 的深度特征作为图像向量。查询图片向量生成后，系统调用 Qdrant "
        "的查询接口返回相似向量及其 Payload。为了兼容不同版本的 `qdrant-client`，服务层优先使用新版 `query_points()` 接口，并保留旧版 `search()` 调用作为兜底。",
    )
    add_formula(doc, "sim(x, y) = cos(x, y) = (x · y) / (||x|| × ||y||)")
    add_screenshot_placeholder(doc, "图5-2 相似检索页面返回相似案例和相似度", "http://127.0.0.1:5000/#similar")
    add_heading(doc, "5.4 分类知识与智能问答实现", 2)
    add_paragraph(
        doc,
        "分类知识模块提供本地规则检索能力，用户可以输入“电池”“塑料瓶”等关键词查询对应类别和投放建议。智能问答模块调用 DeepSeek Chat API，系统提示词限定模型扮演垃圾分类助手，"
        "要求回答尽量围绕废弃物分类、投放方式、注意事项和环保知识展开。该设计能够补充静态知识库覆盖不足的问题。",
    )
    add_screenshot_placeholder(doc, "图5-3 分类知识检索页面，搜索电池或塑料瓶", "http://127.0.0.1:5000/#search")
    add_screenshot_placeholder(doc, "图5-4 智能问答页面，提问过期药品属于什么垃圾", "http://127.0.0.1:5000/#chat")
    add_heading(doc, "5.5 图片理解模块实现", 2)
    add_paragraph(
        doc,
        "图片理解模块面向单纯分类模型难以解释的复杂场景。前端在用户选择图片后立即显示预览，用户确认后上传给 Flask 后端。后端将图片编码为 Base64 data URL，"
        "并调用 Kimi OpenAI 兼容多模态 Chat Completions 接口，请求模型分析图片中可能出现的物品、材质、污染状态和垃圾分类建议。该模块能够提高系统对混合物品和复杂背景图片的解释能力。",
    )
    add_screenshot_placeholder(doc, "图5-5 图片理解页面选择图片后显示预览和 Kimi 分析结果", "http://127.0.0.1:5000/#understand")
    add_heading(doc, "5.6 历史记录模块实现", 2)
    add_paragraph(
        doc,
        "历史记录模块用于展示用户近期识别和检索行为。前端通过 `/api/history` 获取分页数据，并在列表中显示图片预览、识别类别、置信度和创建时间。"
        "为了保证图片预览安全，后端提供 `/media/<path>` 路由，并限制只能读取 `uploads`、`data/raw` 和 `data/processed` 等允许目录下的图片。",
    )
    add_screenshot_placeholder(doc, "图5-6 历史记录页面展示图片、类别和时间", "http://127.0.0.1:5000/#history")
    add_heading(doc, "5.7 前端界面实现", 2)
    add_paragraph(
        doc,
        "系统前端由早期静态 HTML 页面升级为 Vue 3 单页应用，并采用深色侧边栏、浅色工作台、玻璃质感顶栏、多色统计卡和真实废弃物样例图组成现代化管理后台风格。"
        "当前运行版本已根据需求移除登录、注册和用户管理入口，知识测试模块暂时隐藏，仅保留后端接口和部分页面代码以便后续恢复。",
    )

    add_heading(doc, "6 系统测试与结果分析", 1)
    add_heading(doc, "6.1 测试环境", 2)
    add_table(
        doc,
        "表6-1 系统测试环境",
        ["项目", "配置"],
        [
            ["本地系统", "Windows 本地开发环境"],
            ["后端环境", "Python 虚拟环境、Flask"],
            ["前端环境", "Vue 3 全局运行时、浏览器"],
            ["向量数据库", "Docker Qdrant，宿主机端口 63330"],
            ["训练环境", "Slurm GPU 节点，NVIDIA L40S"],
            ["测试命令", ".venv\\Scripts\\python.exe -m pytest -q"],
        ],
    )
    add_heading(doc, "6.2 功能测试", 2)
    add_table(
        doc,
        "表6-2 主要功能测试用例",
        ["编号", "测试内容", "预期结果", "测试结论"],
        [
            ["T01", "访问首页", "页面正常加载并显示核心模块", "通过"],
            ["T02", "上传合法图片识别", "返回类别和置信度", "通过"],
            ["T03", "上传非法格式文件", "返回格式错误提示", "通过"],
            ["T04", "上传超过 16MB 文件", "返回大小限制提示", "通过"],
            ["T05", "相似检索", "返回相似图片及相似度", "通过"],
            ["T06", "分类知识检索", "返回匹配知识内容", "通过"],
            ["T07", "DeepSeek 问答缺配置", "返回配置缺失提示", "通过"],
            ["T08", "Kimi 图片理解缺配置", "返回配置缺失提示", "通过"],
            ["T09", "历史记录分页", "返回分页数据", "通过"],
            ["T10", "Qdrant 客户端兼容", "新版接口可正常调用", "通过"],
        ],
    )
    add_screenshot_placeholder(doc, "图6-1 pytest 自动化测试结果，显示 11 passed", "本地终端运行 .venv\\Scripts\\python.exe -m pytest -q 后截图")
    add_heading(doc, "6.3 模型测试结果", 2)
    add_paragraph(
        doc,
        "ResNet50 模型训练 10 轮后，最佳验证准确率为 0.9821，测试集准确率为 0.9830。分类报告显示，四类垃圾的宏平均 precision、recall 和 f1-score 均达到 0.97，"
        "加权平均 f1-score 为 0.98，说明模型整体分类性能较好。其他垃圾类别样本数量相对较少，precision 为 0.93，后续可通过扩充数据和更精细的数据增强进一步提升稳定性。",
    )
    add_table(
        doc,
        "表6-3 ResNet50 分类报告",
        ["类别", "Precision", "Recall", "F1-score", "Support"],
        [
            ["hazardous", "0.99", "0.98", "0.98", "142"],
            ["kitchen", "0.98", "0.96", "0.97", "134"],
            ["other", "0.93", "0.97", "0.95", "97"],
            ["recyclable", "0.99", "0.99", "0.99", "860"],
            ["accuracy", "-", "-", "0.98", "1233"],
            ["macro avg", "0.97", "0.97", "0.97", "1233"],
            ["weighted avg", "0.98", "0.98", "0.98", "1233"],
        ],
    )
    add_screenshot_placeholder(doc, "图6-2 远端训练日志或分类报告文件截图", "models/classification_report.txt 或 outputs/slurm/resnet50_33280811.out")
    add_heading(doc, "6.4 相似检索测试结果", 2)
    add_paragraph(
        doc,
        "Qdrant 向量索引已在本地 Docker 环境构建完成。由于 Windows 系统保留了默认端口所在的 6256-6355 端口段，本地宿主机使用 63330/63340 映射容器内 6333/6334。"
        "`waste_images` 集合共写入 8213 条图片向量，向量维度为 2048，距离度量为 Cosine，集合状态为 green。使用 `battery__battery1.jpg` 验证相似检索时，首次请求约 14.01 秒，"
        "主要耗时来自模型冷启动；热启动后两次请求约 893.19 毫秒和 891.63 毫秒，返回最高相似度 1.0 的结果。",
    )
    add_screenshot_placeholder(doc, "图6-3 Qdrant 集合状态或相似检索接口响应截图", "http://localhost:63330/dashboard 或相似检索页面结果")
    add_heading(doc, "6.5 测试结果分析", 2)
    add_paragraph(
        doc,
        "从功能测试结果看，系统主要业务流程均能正常运行，异常输入能够被识别并返回友好提示。从模型测试结果看，ResNet50 对可回收物、有害垃圾和厨余垃圾具有较高识别能力，"
        "其他垃圾类别受样本数量和类别边界影响，仍存在进一步优化空间。从性能结果看，模型首次加载会造成明显冷启动耗时，后续请求速度明显提升，因此正式部署时可考虑服务启动时预加载模型。",
    )

    add_heading(doc, "7 总结与展望", 1)
    add_heading(doc, "7.1 总结", 2)
    add_paragraph(
        doc,
        "本文完成了一套基于深度学习的废弃物数据管理系统的设计与实现。系统以 ResNet50 为核心识别模型，结合 Flask、Vue、SQLite 和 Qdrant 构建完整 Web 应用，"
        "并接入 DeepSeek 和 Kimi 提供智能问答与图片理解能力。经过数据整理、模型训练、索引构建和功能测试，系统能够完成垃圾图片识别、相似案例检索、分类知识查询、"
        "智能交流、图片理解和历史记录管理等功能，满足毕业设计任务书提出的主要要求。",
    )
    add_heading(doc, "7.2 展望", 2)
    add_paragraph(
        doc,
        "后续工作可以从以下方面继续优化：第一，扩充数据集规模，尤其增加其他垃圾和复杂混合场景图片，提高模型泛化能力；第二，尝试 EfficientNet、ConvNeXt 或轻量化模型，"
        "在准确率和推理速度之间取得更好平衡；第三，完善 Qdrant 数据管理后台，支持索引更新、重复图片检查和检索结果反馈；第四，恢复并完善知识测试模块，形成学习闭环；"
        "第五，在部署层面引入模型预热、缓存和异步任务机制，进一步提升系统响应速度和稳定性。",
    )

    add_heading(doc, "参考文献", 1)
    references = [
        "[1] He K, Zhang X, Ren S, et al. Deep Residual Learning for Image Recognition[C]. IEEE Conference on Computer Vision and Pattern Recognition, 2016.",
        "[2] Paszke A, Gross S, Massa F, et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library[C]. NeurIPS, 2019.",
        "[3] Pallets Projects. Flask Documentation[EB/OL]. https://flask.palletsprojects.com/.",
        "[4] Qdrant. Qdrant Vector Database Documentation[EB/OL]. https://qdrant.tech/documentation/.",
        "[5] Vue.js Team. Vue 3 Documentation[EB/OL]. https://vuejs.org/.",
        "[6] DeepSeek. DeepSeek API Documentation[EB/OL]. https://api-docs.deepseek.com/.",
        "[7] Moonshot AI. Kimi API Documentation[EB/OL]. https://platform.moonshot.cn/docs/.",
        "[8] Mohamed M. Garbage Classification (12 classes)[DB/OL]. Kaggle, 2021. https://www.kaggle.com/datasets/mostafaabla/garbage-classification.",
        "[9] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016.",
        "[10] Krizhevsky A, Sutskever I, Hinton G E. ImageNet Classification with Deep Convolutional Neural Networks[J]. Communications of the ACM, 2017.",
        "[11] Johnson J, Douze M, Jégou H. Billion-scale similarity search with GPUs[J]. IEEE Transactions on Big Data, 2019.",
    ]
    for item in references:
        add_paragraph(doc, item, indent=False)

    add_heading(doc, "致谢", 1)
    add_paragraph(
        doc,
        "在本课题完成过程中，感谢指导教师在选题、系统设计、论文结构和测试验证等方面给予的指导和帮助。感谢学院提供的学习环境和实践条件，"
        "也感谢开源社区提供的 Flask、Vue、PyTorch、Qdrant 等优秀工具。通过本次毕业设计，我对深度学习模型训练、Web 系统开发、向量检索和智能接口集成有了更加系统的理解，"
        "也进一步提升了独立分析问题和工程实现的能力。",
    )


def normalize_sections(doc: Document) -> None:
    """统一页面边距，减少模板原有多节页面设置带来的排版不一致。"""

    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def main() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"未找到模板文件：{TEMPLATE}")

    doc = Document(TEMPLATE)
    clear_document_body(doc)
    configure_styles(doc)
    normalize_sections(doc)
    add_cover(doc)
    add_abstracts(doc)
    add_toc_placeholder(doc)
    add_body(doc)
    doc.save(OUTPUT)
    print(f"已生成：{OUTPUT}")


if __name__ == "__main__":
    main()
